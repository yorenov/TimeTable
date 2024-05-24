import datetime

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm, Pt
from docx.styles import styles
from docx.styles.style import ParagraphStyle
from datetime import datetime
from cell_tools import *
import db


def create_doc():
    return Document()


def save_doc(document, filename):
    set_margins(document.sections, (Cm(1), Cm(0.5), Cm(0.5), Cm(0.5)))
    document.save("./расписание/" + filename)


def create_word(document=None, date_list=None, group_name='ИСП-23-19', timetable=None):
    if timetable is None:
        timetable = {}
    if date_list is None:
        date_list = ['01.01', '02.01', '03.01', '04.01', '05.01', '06.01']
    if document is None:
        document = Document()
    par0 = document.add_paragraph()
    r1 = par0.add_run('Прием: ')
    r2 = par0.add_run(f"{'_' * ((32 - len(group_name)) // 2)}{group_name}{'_' * ((32 - len(group_name)) // 2)}")
    r1.font.size = Pt(14)
    r2.font.size = Pt(14)
    r2.bold = True
    r2.underline = True

    par1 = document.add_paragraph()
    par1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par1.add_run('РАСПИСАНИЕ ЗАНЯТИЙ')
    par1.runs[0].bold = True
    par1.runs[0].italic = True
    par1.runs[0].font.size = Pt(28)

    table = document.add_table(rows=1, cols=6) # , style='Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # table.autofit = True
    # table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    headers = ['', 'Дата', 'Время', 'Дисциплины', 'Преподаватель', 'Аудитория']
    for i, k in enumerate(hdr_cells):
        k.text = headers[i]
        set_cell_border(k, bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"})
        set_cell_border(k, top={"sz": 12, "val": "single", "color": "#000000", "space": "0"})

        k.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        k.paragraphs[0].runs[0].bold = True
        if i != 0 and i != 5:
            set_cell_border(k, start={"sz": 6, "val": "single", "color": "#000000", "space": "0"}, end={"sz": 6, "val": "single", "color": "#000000", "space": "0"})

    for i, day in enumerate(['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']):
        row_cells = table.add_row().cells
        row_cells[0].text = day
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_vertical_cell_direction(row_cells[0], 'btLr')

        row_cells[1].text = date_list[i]
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_cell_border(row_cells[0], top={"sz": 6, "val": "single", "color": "#000000", "space": "0"}, bottom={"sz": 6, "val": "single", "color": "#000000", "space": "0"})
        set_cell_border(row_cells[1], start={"sz": 6, "val": "single", "color": "#000000", "space": "0"}, top={"sz": 6, "val": "single", "color": "#000000", "space": "0"}, bottom={"sz": 6, "val": "single", "color": "#000000", "space": "0"})

        row_cells[2].merge(row_cells[3])
        row_cells[2].merge(row_cells[4])
        row_cells[2].merge(row_cells[5])

        items = [
            [x.time for x in timetable[str(i)]],
            [x.subject for x in timetable[str(i)]],
            [x.teacher for x in timetable[str(i)]],
            [x.audience for x in timetable[str(i)]]
        ]

        micro_table(row_cells[2], items)
        set_cell_border(row_cells[2], top={"sz": 12, "val": "single", "color": "#000000", "space": "0"}, bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"})
    set_col_widths(table, (Cm(1), Cm(2), Cm(4.5), Cm(4.5), Cm(4.5), Cm(4.5)))
    r_last = document.add_paragraph()
    r_last.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d = datetime.now().strftime("%d.%m %H:%M")
    r_last.text = f"Сгенерировано программой TimeTable by yorenov в {d}"
    document.add_page_break()

    return document


if __name__ == '__main__':
    doc = create_doc()
    doc = create_word(doc, ['1', '2', '3', '4', '5', '6'], group_name='Группа 1')
