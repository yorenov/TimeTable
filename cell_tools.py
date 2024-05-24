from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
from docx.table import _Cell, Table

TABLE_IN_TABLE_ROW_HEIGHT = Cm(0.5)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_margins(sections, margins):
    for section in sections:
        section.top_margin = margins[0]
        section.bottom_margin = margins[1]
        section.left_margin = margins[2]
        section.right_margin = margins[3]


def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def set_table_border(table: Table, **kwargs):
    """
    Sets table border
    Usage:

    set_table_border(
        table,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tbl  = table._tbl
    tblPr = tbl.tblPr

    # check for tag existnace, if none found, then create one
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tblBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tblBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def micro_table(cell: _Cell, items):
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    r = cell.add_table(rows=max([len(x) for x in items]), cols=len(items))
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    for i, r_row in enumerate(r.rows):
        r_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        max_lines = 0
        for cc in items:
            max_lines = max(max_lines, len(cc[i]) // 12)
        r_row.height = TABLE_IN_TABLE_ROW_HEIGHT * max_lines
        for i_c, r_cell in enumerate(r_row.cells):
            r_cell._element.clear_content()
            set_cell_margins(r_cell, top=0, start=0, bottom=0, end=0)
            r_cell.add_paragraph().add_run(items[i_c][i])
            r_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i_c == 0:
                r_cell.width = Cm(4.45)
            elif i_c == 1:
                r_cell.width = Cm(4.48)
            else:
                r_cell.width = Cm(4.50)
            set_cell_border(r_cell, start={"sz": 6, "val": "single", "color": "#000000", "space": "0"})
            if i != len(r.rows) - 1:
                set_cell_border(r_cell, bottom={"sz": 6, "val": "single", "color": "#000000", "space": "0"})